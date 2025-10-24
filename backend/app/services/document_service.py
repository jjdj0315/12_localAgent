"""
Document service for file upload, validation, and text extraction.
"""
import os
import magic
from pathlib import Path
from typing import Optional, Tuple
from uuid import UUID, uuid4

import pdfplumber
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.document import Document
from app.models.conversation_document import ConversationDocument


class DocumentService:
    """Service for document operations"""

    # Configuration
    UPLOAD_DIR = Path("uploads")
    MAX_FILE_SIZE = 52428800  # 50MB in bytes

    # Allowed MIME types and extensions
    ALLOWED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
    }

    @staticmethod
    def ensure_upload_dir():
        """Ensure upload directory exists"""
        DocumentService.UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

    @staticmethod
    async def validate_file(
        file_content: bytes, filename: str
    ) -> Tuple[bool, str, Optional[str]]:
        """
        Validate file type and size using magic numbers.

        Args:
            file_content: File content bytes
            filename: Original filename

        Returns:
            Tuple of (is_valid, error_message, file_type)
        """
        # Check file size
        file_size = len(file_content)
        if file_size > DocumentService.MAX_FILE_SIZE:
            return (
                False,
                "파일이 너무 큽니다. 50MB 이하의 파일을 업로드해주세요.",
                None,
            )

        if file_size == 0:
            return False, "파일이 비어있습니다.", None

        # Check file type with magic number (first 2048 bytes)
        try:
            mime_type = magic.from_buffer(file_content[:2048], mime=True)
        except Exception:
            return False, "파일 형식을 확인할 수 없습니다.", None

        if mime_type not in DocumentService.ALLOWED_TYPES:
            return (
                False,
                "지원하지 않는 파일 형식입니다. PDF, DOCX, TXT 파일만 업로드 가능합니다.",
                None,
            )

        file_type = DocumentService.ALLOWED_TYPES[mime_type]
        return True, "", file_type

    @staticmethod
    async def extract_text_from_pdf(file_path: Path) -> str:
        """
        Extract text from PDF file using pdfplumber.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        try:
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                return "\n\n".join(pages_text)
        except Exception as e:
            raise ValueError(f"PDF 파일 처리 중 오류가 발생했습니다: {str(e)}")

    @staticmethod
    async def extract_text_from_docx(file_path: Path) -> str:
        """
        Extract text from DOCX file using python-docx.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            doc = DocxDocument(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise ValueError(f"DOCX 파일 처리 중 오류가 발생했습니다: {str(e)}")

    @staticmethod
    async def extract_text_from_txt(file_path: Path) -> str:
        """
        Extract text from TXT file.

        Args:
            file_path: Path to TXT file

        Returns:
            Extracted text content
        """
        try:
            # Try UTF-8 first, then fall back to other encodings
            encodings = ["utf-8", "utf-8-sig", "cp949", "euc-kr"]
            for encoding in encodings:
                try:
                    return file_path.read_text(encoding=encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("텍스트 파일 인코딩을 인식할 수 없습니다.")
        except Exception as e:
            raise ValueError(f"TXT 파일 처리 중 오류가 발생했습니다: {str(e)}")

    @staticmethod
    async def extract_text(file_path: Path, file_type: str) -> str:
        """
        Extract text from document based on file type.

        Args:
            file_path: Path to document file
            file_type: File type (pdf, docx, txt)

        Returns:
            Extracted text content
        """
        if file_type == "pdf":
            return await DocumentService.extract_text_from_pdf(file_path)
        elif file_type == "docx":
            return await DocumentService.extract_text_from_docx(file_path)
        elif file_type == "txt":
            return await DocumentService.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"지원하지 않는 파일 형식: {file_type}")

    @staticmethod
    async def save_document(
        db: AsyncSession,
        user_id: UUID,
        filename: str,
        file_content: bytes,
        file_type: str,
    ) -> Document:
        """
        Save document to filesystem and database.

        Args:
            db: Database session
            user_id: User ID
            filename: Original filename
            file_content: File content bytes
            file_type: File type (pdf, docx, txt)

        Returns:
            Created Document instance
        """
        # Ensure upload directory exists
        DocumentService.ensure_upload_dir()

        # Create user-specific directory
        user_dir = DocumentService.UPLOAD_DIR / str(user_id)
        user_dir.mkdir(parents=True, exist_ok=True)

        # Generate unique document ID and file path
        doc_id = uuid4()
        file_path = user_dir / f"{doc_id}.{file_type}"

        # Save file to disk
        file_path.write_bytes(file_content)

        # Extract text from document
        try:
            extracted_text = await DocumentService.extract_text(file_path, file_type)
        except Exception as e:
            # Clean up file if text extraction fails
            file_path.unlink(missing_ok=True)
            raise e

        # Create database record
        document = Document(
            id=doc_id,
            user_id=user_id,
            filename=filename,
            file_path=str(file_path),
            file_type=file_type,
            file_size=len(file_content),
            extracted_text=extracted_text,
        )

        db.add(document)
        await db.commit()
        await db.refresh(document)

        return document

    @staticmethod
    async def list_documents(
        db: AsyncSession, user_id: UUID, page: int = 1, page_size: int = 20
    ) -> Tuple[list[Document], int]:
        """
        List user's documents with pagination.

        Args:
            db: Database session
            user_id: User ID
            page: Page number (1-indexed)
            page_size: Items per page

        Returns:
            Tuple of (documents list, total count)
        """
        # Get total count
        from sqlalchemy import func

        count_query = select(func.count()).where(Document.user_id == user_id)
        total_result = await db.execute(count_query)
        total = total_result.scalar() or 0

        # Get documents with pagination
        offset = (page - 1) * page_size
        query = (
            select(Document)
            .where(Document.user_id == user_id)
            .order_by(Document.uploaded_at.desc())
            .offset(offset)
            .limit(page_size)
        )

        result = await db.execute(query)
        documents = result.scalars().all()

        return list(documents), total

    @staticmethod
    async def get_document(
        db: AsyncSession, document_id: UUID, user_id: UUID
    ) -> Optional[Document]:
        """
        Get document by ID with user authorization.

        Args:
            db: Database session
            document_id: Document ID
            user_id: User ID (for authorization)

        Returns:
            Document or None if not found
        """
        from sqlalchemy import and_

        query = select(Document).where(
            and_(
                Document.id == document_id,
                Document.user_id == user_id,
            )
        )

        result = await db.execute(query)
        return result.scalar_one_or_none()

    @staticmethod
    async def delete_document(
        db: AsyncSession, document_id: UUID, user_id: UUID
    ) -> bool:
        """
        Delete document from database and filesystem.

        Args:
            db: Database session
            document_id: Document ID
            user_id: User ID (for authorization)

        Returns:
            True if deleted, False if not found
        """
        document = await DocumentService.get_document(db, document_id, user_id)

        if not document:
            return False

        # Delete file from filesystem
        try:
            file_path = Path(document.file_path)
            if file_path.exists():
                file_path.unlink()
        except Exception:
            # Continue even if file deletion fails
            pass

        # Delete from database
        await db.delete(document)
        await db.commit()

        return True

    @staticmethod
    async def attach_document_to_conversation(
        db: AsyncSession, conversation_id: UUID, document_id: UUID
    ) -> None:
        """
        Attach document to conversation.

        Args:
            db: Database session
            conversation_id: Conversation ID
            document_id: Document ID
        """
        # Check if already attached
        query = select(ConversationDocument).where(
            ConversationDocument.conversation_id == conversation_id,
            ConversationDocument.document_id == document_id,
        )
        result = await db.execute(query)
        existing = result.scalar_one_or_none()

        if not existing:
            link = ConversationDocument(
                conversation_id=conversation_id,
                document_id=document_id,
            )
            db.add(link)
            await db.commit()

    @staticmethod
    async def get_conversation_documents(
        db: AsyncSession, conversation_id: UUID
    ) -> list[Document]:
        """
        Get all documents attached to a conversation.

        Args:
            db: Database session
            conversation_id: Conversation ID

        Returns:
            List of documents
        """
        query = (
            select(Document)
            .join(ConversationDocument)
            .where(ConversationDocument.conversation_id == conversation_id)
        )

        result = await db.execute(query)
        return list(result.scalars().all())


# Singleton instance
document_service = DocumentService()
